import sqlite3
from flask import Flask, render_template, redirect, request, flash, send_from_directory
from werkzeug.exceptions import abort
import os
from pprint import pprint
from docx import Document

import datetime

from init_db import inserts
from init_db import updates


app = Flask(__name__)
app.config['SECRET_KEY'] = b'my)secret)key'
UPLOAD_FOLDER_CONTRACTS = 'contracts'
UPLOAD_FOLDER_REPORTS = 'reports'
app.config['UPLOAD_FOLDER_CONTRACTS'] = UPLOAD_FOLDER_CONTRACTS
app.config['UPLOAD_FOLDER_REPORTS'] = UPLOAD_FOLDER_REPORTS

def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/')
def index():
    return redirect("/contracts")


###############!                         ###############
###############!                         ###############
###############! Договоры (НЕ КОНТРАКТЫ) ###############
###############!                         ###############
###############!                         ###############

@app.route('/contracts')
def contracts():
    """ Страница-список - получение всех контрактов """
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM contracts, employees, suppliers
        WHERE suppliers.id = contracts.supplier_id 
        and employees.id = contracts.employee_id 
        ORDER BY contracts.date"""
    ).fetchall()
    conn.close()
    return render_template('contract/contracts.html', 
                           contracts=res)


# используется в contract(contract_id)
def get_contract(item_id):
    """ Получение одного контракта из БД """
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, employees, suppliers
        WHERE contracts.employee_id = employees.id 
        and contracts.supplier_id = suppliers.id 
        and contracts.id = ?
    """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    """ Страница-карточка - 1 контракт """
    pos = get_contract(contract_id)
    emp = get_employee(pos['employee_id'])
    sup = get_supplier(pos['supplier_id'])
    return render_template('contract/contract.html', 
                           contract=pos, 
                           employee=emp,
                           supplier=sup)


@app.route('/new_contract', methods=('GET', 'POST'))
def new_contract():
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            deal_type = request.form['deal_type']
            description = request.form['description']
            status = int(request.form['status'])
            total_cost = int(request.form['total_cost'])
            supplier = int(request.form.get('supplier'))
            employee = int(request.form.get('employee'))
            bundle = [
                number, date, deal_type, description,
                status, total_cost, supplier, employee
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all((number, date, deal_type, total_cost, supplier, employee)):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                inserts["contracts"], 
                bundle
            )
            conn.commit()
            new_contract_id = cursor.lastrowid
            
            conn.close()
            return redirect(f'/contract/{new_contract_id}')
    # отрисовка формы        
    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM employees""").fetchall()
    pos2 = conn.execute("""SELECT * FROM suppliers""").fetchall()
    conn.close()
    return render_template('contract/new_contract.html', employees=pos1, suppliers=pos2)


@app.route('/edit_contract', methods=('GET', 'POST'))
def edit_contract():
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            deal_type = request.form['deal_type']
            description = request.form['description']
            status = int(request.form['status'])
            total_cost = int(request.form['total_cost'])
            supplier = int(request.form.get('supplier'))
            employee = int(request.form.get('employee'))
            bundle = [
                number, date, deal_type, description,
                status, total_cost, supplier, employee
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all((number, date, deal_type, total_cost, supplier, employee)):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            cursor = conn.cursor()
            contract_id = request.args["contract_id"]
            buffer = bundle[::]
            buffer.append(contract_id)
            cursor.execute(
                updates["contracts"], 
                buffer
            )
            conn.commit()
            
            conn.close()
            return redirect(f'/contract/{contract_id}')
    # отрисовка формы        
    conn = get_db_connection()
    contract = get_contract(request.args["contract_id"])
    pos1 = conn.execute("""SELECT * FROM employees""").fetchall()
    pos2 = conn.execute("""SELECT * FROM suppliers""").fetchall()
    conn.close()
    return render_template('contract/edit_contract.html', contract=contract, employees=pos1, suppliers=pos2)


@app.route("/delete_contract/<int:contract_id>", methods=('GET', "POST", "DELETE"))
def delete_contract(contract_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """DELETE FROM contracts WHERE id = ?""", 
        (contract_id,)
    )
    conn.commit()
    conn.close()
    return redirect("/contracts")


#webpage
@app.route('/generate_contract', methods=('GET', 'POST'))
def generate_contract():
    """ Страница генерации договора """
    # переменные шаблона
    id = int(request.args.get('id_contract'))
    conn = get_db_connection()
    contract = get_contract(id)
    supplier = get_supplier(contract['supplier_id'])
    employee = get_employee(contract['employee_id'])
    conn.close()
    # manual typing-in params
    contract_params = {
        "CONTRACT_DEAL_OTHER_INFO" : "Дополнительные сведения об услугах",
        "SUPPLIER_ATTORNEY_POWER" : "Юридическая сила поставщика",
        "EMPLOYEE_ATTORNEY_POWER" : "Юридическая сила сотрудника",
        "GOODS_NAME" : "Наименование товара",
        "GOODS_ADDRESS" : "Адрес заказчика",
        "GOODS_AMOUNT" : "Количество товаров",
        "GOODS_PRICE" : "Цена за штуку",
        "GOODS_UNIT" : "Единица измерения",
        "CONTRACT_DEAL_LASTDATE" : "Срок поставки (до [...])",
        "CONTRACT_DEAL_PAYMENT_INFO" : "Срок поставки (в течение [...])",
        "CONTRACT_CANCEL_AFTER" : "Количество дней для отмены сделки после получения товара",
        "CONTRACT_CANCEL_BEFORE" : "Количество дней для отмены сделки без получения товара",
        "CONTRACT_LASTDATE" : "Дата завершения действия договора",        
    }
    # automatic params from db
    contract_params_auto = {
        'CONTRACT_NUMBER': ['номер договора', contract['number']],
        'CONTRACT_DATE': ['дата подписания договора', contract['date']],
        'CONTRACT_DEAL_TYPE': ['тип договора', contract['deal_type']],
        'CONTRACT_TOTAL_PRICE': ['общая сумма', contract['total_cost']],
        'EMPLOYEE_FULLNAME': ['ФИО', employee['name']],
        'EMPLOYEE_PHONE_NUMBER': ['номер телефона сотрудника', employee['phone_number']],
        'EMPLOYEE_POSITION': ['должность', employee['position']],
        'SUPPLIER_COMPANY_NAME': ['полное юр. лицо компании', supplier['company_name']],
        'SUPPLIER_OWNER_NAME': ['владелец компании поставщика', supplier['owner_name']],
        'SUPPLIER_POSTAL_CODE': ['полное юр. лицо компании', supplier['postal_code']],
        'SUPPLIER_ADDRESS': ['полное юр. лицо компании', supplier['address']],
        'SUPPLIER_PHONE_NUMBER': ['полное юр. лицо компании', supplier['phone_number']],
    }
    if request.method == 'POST':
        # создание нового документа
        result_params =  request.form.to_dict()
        create_contract(id, result_params)
        return redirect(f'/contract/{id}')
    # скачивание файла, если он заполнен
    filename = f"Договор {contract['number']} от {contract['date']}.docx"
    if os.path.exists(os.path.join('contracts', filename)):
        return send_from_directory(app.config['UPLOAD_FOLDER_CONTRACTS'], filename, as_attachment=True)
    else:
        # отрисовка формы заполнения
        flash('Договор не сформирован, заполните его')
        return render_template('contract/generate_contract.html', 
                               contract=contract,
                               employee=employee,
                               supplier=supplier,
                               contract_params=contract_params,
                               auto_params=contract_params_auto)


#python-docx
def create_contract(id, contract_params):
    """ Создание нового документа по шаблону """
    template = os.path.join('contracts', 'contract_template.docx')
    result = os.path.join('contracts', f"Договор {contract_params['CONTRACT_NUMBER']} от {contract_params['CONTRACT_DATE']}.docx")
    template_doc = Document(template)
    for key, value in contract_params.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, f'=={key}==', value)
        for table in template_doc.tables:
            replace_text_in_tables(table, f'=={key}==', value)
    template_doc.save(result)


#python-docx
def replace_text(paragraph, key, value):
    """ Работа docx - заполнение параграфов """
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(str(key), str(value))


#python-docx
def replace_text_in_tables(table, key, value):
    """ Работа docx - заполнение таблиц """
    for row in table.rows:
        for cell in row.cells:
            if key in cell.text:
                cell.text = cell.text.replace(key, value)



###############!            ###############
###############!            ###############
###############! Поставщики ###############
###############!            ###############
###############!            ###############

@app.route('/suppliers')
def suppliers():
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM suppliers"""
    ).fetchall()
    conn.close()
    return render_template('supplier/suppliers.html', 
                           suppliers=res)


def get_supplier(id: int):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM suppliers
        WHERE suppliers.id = ?
    """, (id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/supplier/<int:supplier_id>')
def supplier(supplier_id):
    pos = get_supplier(supplier_id)
    return render_template('supplier/supplier.html',
                           supplier=pos)


@app.route('/new_supplier', methods=('GET', 'POST'))
def new_supplier():
    if request.method == 'POST':
        try:
            company_name = request.form['company_name']
            owner_name = request.form['owner_name']
            address = request.form['address']
            postal_code = request.form['postal_code']
            phone_number = request.form['phone_number']
            bundle = [
                company_name,owner_name, address,
                postal_code, phone_number
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all(bundle):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                inserts["suppliers"], 
                bundle
            )
            conn.commit()
            new_supplier_id = cursor.lastrowid
            
            conn.close()
            return redirect(f'/supplier/{new_supplier_id}')
    return render_template('supplier/new_supplier.html')


@app.route('/edit_supplier', methods=('GET', 'POST'))
def edit_supplier():
    if request.method == 'POST':
        try:
            company_name = request.form['company_name']
            owner_name = request.form['owner_name']
            address = request.form['address']
            postal_code = request.form['postal_code']
            phone_number = request.form['phone_number']
            bundle = [
                company_name,owner_name, address,
                postal_code, phone_number
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all(bundle):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            cursor = conn.cursor()
            supplier_id = request.args["supplier_id"]
            buffer = bundle[::]
            buffer.append(supplier_id)
            cursor.execute(
                inserts["suppliers"], 
                buffer
            )
            conn.commit()
            
            conn.close()
            return redirect(f'/supplier/{supplier_id}')
    conn = get_db_connection()
    cur = conn.cursor()
    supplier_id = request.args["supplier_id"]
    sup = cur.execute(
        """SELECT * FROM suppliers WHERE id = ?""", 
        (supplier_id,)).fetchone()
    conn.commit()
    conn.close()
    return render_template('supplier/edit_supplier.html', supplier=sup)


@app.route("/delete_supplier/<int:supplier_id>", methods=('GET', "POST", "DELETE"))
def delete_supplier(supplier_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """DELETE FROM suppliers WHERE id = ?""", 
        (supplier_id,)
    )
    conn.commit()
    conn.close()
    return redirect("/suppliers")


# @app.route()
# def edit_supplier(supplier_id):

###############!         ###############
###############! Сладости ###############
###############!         ###############

@app.route('/sweets')
def sweets():
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM sweets"""
    ).fetchall()
    conn.close()
    return render_template('sweet/sweets.html', 
                           sweets=res)


@app.route('/sweet/<int:sweet_id>')
def sweet(sweet_id: int):
    res = get_sweet(sweet_id)
    supplier_id = res["supplier_id"]
    contract_id = res['contract_id']
    sup = get_supplier(supplier_id)
    try:
        con = get_contract(contract_id)
    except:
        con = ""
    return render_template('sweet/sweet.html',
                           sweet=res,
                           supplier=sup,
                           contract=con)


def get_sweet(id: int):
    conn = get_db_connection()
    res = conn.execute("""
        SELECT * FROM sweets
        WHERE sweets.id = ?
        """, (id,) 
        ).fetchone()
    conn.close()
    if res is None:
        abort(404)
    return res


@app.route('/new_sweet', methods=('GET', 'POST'))
def new_sweet():
    if request.method == 'POST':
        try:
            name = request.form['name'] # text
            manufacturer = request.form['manufacturer'] # text
            country = request.form['country'] # text
            weight = abs(float(request.form['weight']))
            price = abs(int(request.form['price']))
            calories = int(request.form['calories'])
            production_date = request.form['production_date']
            expiration_date = request.form['expiration_date']
            # check for bad value as well
            supplier = int(request.form.get('supplier'))
            contract = int(request.form.get('contract'))
            #get shit together yo
            bundle = [
                name, manufacturer, country,  
                weight, price, calories, 
                production_date, expiration_date, supplier, contract
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all(bundle):
            flash('Не все поля заполнены')
        else:    
            sweet_id = -1
            if (supplier == -1):
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    inserts["sweets_no_contract"],
                    bundle[:-1])
                conn.commit()
                sweet_id = cur.lastrowid
                cur.close()
            else:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    inserts["sweets"],
                    bundle)
                conn.commit()
                sweet_id = cur.lastrowid
                cur.close()
            return redirect(f'/sweet/{sweet_id}')
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM 'suppliers'""").fetchall()
    res2 = conn.execute("""SELECT *, company_name FROM 'contracts', 'suppliers' 
                        WHERE contracts.supplier_id = suppliers.id""").fetchall()
    conn.close()
    return render_template('sweet/new_sweet.html',
                           suppliers=res, contracts=res2)


@app.route('/edit_sweet', methods=('GET', 'POST'))
def edit_sweet():
    if request.method == 'POST':
        try:
            name = request.form['name'] # text
            manufacturer = request.form['manufacturer'] # text
            country = request.form['country'] # text
            weight = abs(float(request.form['weight']))
            price = abs(int(request.form['price']))
            calories = int(request.form['calories'])
            production_date = request.form['production_date']
            expiration_date = request.form['expiration_date']
            # check for bad value as well
            supplier = int(request.form.get('supplier'))
            contract = int(request.form.get('contract'))
            #get shit together yo
            bundle = [
                name, manufacturer, country,  
                weight, price, calories, 
                production_date, expiration_date, supplier, contract
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all(bundle):
            flash('Не все поля заполнены')
        else:    
            sweet_id = request.args['sweet_id']
            if (supplier == -1):
                conn = get_db_connection()
                cur = conn.cursor()
                buffer = bundle[:-1]
                buffer.append(sweet_id)
                cur.execute(
                    updates["sweets_no_contract"],
                    buffer)
                conn.commit()
                cur.close()
            else:
                conn = get_db_connection()
                cur = conn.cursor()
                buffer = bundle[::]
                buffer.append(sweet_id)
                cur.execute(
                    updates["sweets"],
                    buffer)
                conn.commit()
                cur.close()
            return redirect(f'/sweet/{sweet_id}')
    conn = get_db_connection()
    sweet = get_sweet(request.args['sweet_id'])
    res = conn.execute("""SELECT * FROM 'suppliers'""").fetchall()
    res2 = conn.execute("""SELECT *, company_name FROM 'contracts', 'suppliers' 
                        WHERE contracts.supplier_id = suppliers.id""").fetchall()
    conn.close()
    return render_template('sweet/edit_sweet.html',
                           sweet=sweet,
                           suppliers=res, contracts=res2)


@app.route("/delete_sweet/<int:sweet_id>", methods=('GET', "POST", "DELETE"))
def delete_sweet(sweet_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """DELETE FROM sweets WHERE id = ?""", 
        (sweet_id,)
    )
    conn.commit()
    conn.close()
    return redirect("/sweets")

    
###############!           ###############
###############! Виды кофе ###############
###############!           ###############

@app.route('/coffee_varieties')
def coffee_varieties():
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM coffee_varieties"""
    ).fetchall()
    conn.close()
    return render_template('coffee_variety/coffee_varieties.html', 
                           coffee_varieties=res)


@app.route('/coffee_variety/<int:coffee_variety_id>')
def coffee_variety(coffee_variety_id: int):
    res = get_coffee_variety(coffee_variety_id)
    supplier_id = res["supplier_id"]
    contract_id = res["contract_id"]
    sup = get_supplier(supplier_id)
    try:
        con = get_contract(contract_id)
    except:
        con = ""
    return render_template('coffee_variety/coffee_variety.html',
                           coffee_variety=res,
                           supplier=sup,
                           contract=con
                           )


def get_coffee_variety(id: int):
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM coffee_varieties
    WHERE coffee_varieties.id = ?                   
    """, (id,) ).fetchone()
    conn.close()
    if res is None:
        abort(404)
    return res


def plato(x: int | float, 
          top: int | float, 
          bottom: int | float
          ) -> int | float :
    if top < bottom:
        top, bottom = bottom, top
    if x > top: return top
    if x < bottom: return bottom
    return x


@app.route('/new_coffee_variety', methods=('GET', 'POST'))
def new_coffee_variety():
    if request.method == 'POST':
        try:
            name = request.form['name'] # text
            manufacturer = request.form['manufacturer'] # text
            country = request.form['country'] # text
            a = lambda x: int(plato(int(x), 1, 5))
            intensity = a(str(request.form['intensity']).strip())
            acidity   = a(str(request.form['acidity']  ).strip())
            sweetness = a(str(request.form['sweetness']).strip())
            
            roast = request.form['roast'] # text
            dark_or_light_roast = bool(request.form['dark_or_light_roast'])
            weight = abs(float(request.form['weight']))
            price = abs(int(request.form['price']))
            production_date = request.form['production_date']
            expiration_date = request.form['expiration_date']
            # check for bad value as well
            supplier = int(request.form.get('supplier'))
            contract = int(request.form.get('contract'))
            #get shit together yo
            bundle = [
                name, manufacturer, country, intensity, 
                acidity, sweetness, roast, dark_or_light_roast, 
                weight, price, production_date, expiration_date, supplier, contract
            ]
        except ValueError:
            flash('Некорректные значения')
            intensity = -1
            acidity = -1
            sweetness = -1
            roast = ""
            weight = -1.0
            price = -1
            production_date = ""
            expiration_date = ""
            dark_or_light_roast = False
            supplier = -5
            contract = -5
            bundle = [
                name, manufacturer, country, intensity, 
                acidity, sweetness, roast, dark_or_light_roast, 
                weight, price, production_date, expiration_date, supplier, contract
            ]
        if not all(bundle[:7]) and not all(bundle[8:-1]) and dark_or_light_roast != -1:
            flash('Не все поля заполнены')
        elif bundle and len(bundle) != 14:
            flash('Не все поля заполнены')
        else:    
            coffee_id = -1
            if (supplier == -1):
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    inserts["coffee_varieties_no_contract"],
                    bundle[:-1])
                conn.commit()
                coffee_id = cur.lastrowid
                cur.close()
            else:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute(
                    inserts["coffee_varieties"],
                    bundle)
                conn.commit()
                coffee_id = cur.lastrowid
                cur.close()
            return redirect(f'/coffee_variety/{coffee_id}')
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM 'suppliers'""").fetchall()
    res2 = conn.execute("""SELECT *, company_name FROM 'contracts', 'suppliers' 
                        WHERE contracts.supplier_id = suppliers.id""").fetchall()
    conn.close()
    return render_template('coffee_variety/new_coffee_variety.html',
                           suppliers=res, contracts=res2)


@app.route('/edit_coffee_variety', methods=('GET', 'POST'))
def edit_coffee_variety():
    if request.method == 'POST':
        try:
            name = request.form['name'] # text
            manufacturer = request.form['manufacturer'] # text
            country = request.form['country'] # text
            a = lambda x: int(plato(int(x), 1, 5))
            intensity = a(str(request.form['intensity']).strip())
            acidity   = a(str(request.form['acidity']  ).strip())
            sweetness = a(str(request.form['sweetness']).strip())
            
            roast = request.form['roast'] # text
            dark_or_light_roast = bool(request.form['dark_or_light_roast'])
            weight = abs(float(request.form['weight']))
            price = abs(int(request.form['price']))
            production_date = request.form['production_date']
            expiration_date = request.form['expiration_date']
            
            # check for bad value as well
            supplier = int(request.form.get('supplier'))
            contract = int(request.form.get('contract'))
            #get shit together yo
            bundle = [
                name, manufacturer, country, intensity, 
                acidity, sweetness, roast, dark_or_light_roast, 
                weight, price, production_date, expiration_date, supplier, contract
            ]
        except ValueError:
            flash('Некорректные значения')
            intensity = -1
            acidity = -1
            sweetness = -1
            roast = ""
            weight = -1.0
            price = -1
            production_date = ""
            expiration_date = ""
            dark_or_light_roast = False
            supplier = -5
            contract = -5
            bundle = [
                name, manufacturer, country, intensity, 
                acidity, sweetness, roast, dark_or_light_roast, 
                weight, price, production_date, expiration_date, supplier, contract
            ]
        if not all(bundle[:7]) and not all(bundle[8:-1]) and dark_or_light_roast != -1:
            flash('Не все поля заполнены')
        elif bundle and len(bundle) != 14:
            flash('Не все поля заполнены')
        else:    
            coffee_id = request.args['coffee_variety_id']
            if (supplier == -1):
                conn = get_db_connection()
                cur = conn.cursor()
                buffer = bundle[:-1]
                buffer.append(coffee_id)
                cur.execute(
                    updates["coffee_varieties_no_contract"],
                    buffer)
                conn.commit()
                cur.close()
            else:
                conn = get_db_connection()
                cur = conn.cursor()
                buffer = bundle[::]
                buffer.append(coffee_id)
                cur.execute(
                    updates["coffee_varieties"],
                    buffer)
                conn.commit()
                cur.close()
            return redirect(f'/coffee_variety/{coffee_id}')
    conn = get_db_connection()
    coffee = get_coffee_variety(request.args['coffee_variety_id'])
    res = conn.execute("""SELECT * FROM 'suppliers'""").fetchall()
    res2 = conn.execute("""SELECT *, company_name FROM 'contracts', 'suppliers' 
                        WHERE contracts.supplier_id = suppliers.id""").fetchall()
    conn.close()
    return render_template('coffee_variety/edit_coffee_variety.html',
                           coffee_variety=coffee,
                           suppliers=res, contracts=res2)


@app.route("/delete_coffee_variety/<int:coffee_variety_id>", methods=('GET', "POST", "DELETE"))
def delete_coffee_variety(coffee_variety_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """DELETE FROM coffee_varieties WHERE id = ?""", 
        (coffee_variety_id,)
    )
    conn.commit()
    conn.close()
    return redirect("/coffee_varieties")


###############!        ###############
###############! Отчёты ###############
###############!        ###############

@app.route('/reports')
def reports():
    conn = get_db_connection()
    rep = conn.execute("""SELECT * FROM reports, employees
        WHERE reports.employee_id = employees.id
        ORDER BY reports.number
        """).fetchall()
    conn.close()
    return render_template('report/reports.html',
                           reports=rep)


def get_report(id: int):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM reports
        WHERE reports.id = ?
    """, (id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/report/<int:report_id>')
def report(report_id):
    pos = get_report(report_id)
    empl_id = pos["employee_id"]
    employee = get_employee(empl_id)
    return render_template('report/report.html', report=pos, employee=employee)


@app.route('/new_report', methods=('GET', 'POST'))
def new_report():
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            report_type = request.form['report_type']
            description = request.form['description']
            money_earned = int(request.form['money_earned'])
            money_spent = int(request.form['money_spent'])
            employee = int(request.form.get('employee'))
            bundle = [
                number, date, report_type, description,
                money_earned, money_spent, employee
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all(bundle):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute(
                inserts["reports"], 
                bundle
            )
            conn.commit()
            new_report_id = cursor.lastrowid
            conn.close()
            return redirect(f'/report/{new_report_id}')
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM 'employees'""").fetchall()
    conn.close()
    return render_template('report/new_report.html',
                           employees=res)


@app.route('/edit_report', methods=('GET', 'POST'))
def edit_report():
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            report_type = request.form['report_type']
            description = request.form['description']
            money_earned = int(request.form['money_earned'])
            money_spent = int(request.form['money_spent'])
            employee = int(request.form.get('employee'))
            bundle = [
                number, date, report_type, description,
                money_earned, money_spent, employee
            ]
        except ValueError:
            flash('Некорректные значения')
        if not all(bundle):
            flash('Не все поля заполнены')
        else:
            report_id = request.args["report_id"]
            conn = get_db_connection()
            cursor = conn.cursor()
            buffer = bundle[::]
            buffer.append(report_id)
            cursor.execute(
                updates["reports"], 
                buffer
            )
            conn.commit()
            conn.close()
            return redirect(f'/report/{report_id}')
    conn = get_db_connection()
    report = get_report(request.args["report_id"])
    res = conn.execute("""SELECT * FROM 'employees'""").fetchall()
    conn.close()
    return render_template('report/edit_report.html',
                           employees=res, report=report)


@app.route('/generate_report', methods=('GET', 'POST'))
def generate_report():
    """ Страница генерации отчёта """
    # переменные шаблона
    id = int(request.args.get('id_report'))
    conn = get_db_connection()
    report = get_report(id)
    employee = get_employee(report['employee_id'])
    goods1 = conn.execute("""
        SELECT name, price, company_name 
        FROM coffee_varieties, suppliers 
        WHERE coffee_varieties.supplier_id = suppliers.id
        """).fetchall()
    goods2 = conn.execute("""
        SELECT name, price, company_name 
        FROM sweets, suppliers 
        WHERE sweets.supplier_id = suppliers.id
        """).fetchall()
    goods = []
    for i in goods1:
        line = []
        for stuff in i:
            line.append(stuff)
        goods.append(line)
    for i in goods2:
        line = []
        for stuff in i:
            line.append(stuff)
        goods.append(line)

    conn.close()
    
    # manual typing-in params
    report_params = {
        "REPORT_START_DATE" : "Дата начала отсчёта для отчёта",
        "REPORT_FINISH_DATE" : "Дата конца отсчёта для отчёта",
        "REPORT_TOTAL_COST" : "Общая стоимость заказов",
    }
    # automatic params from db
    report_params_auto = {
        'REPORT_NUMBER': ['номер договора', report['number']],
        'REPORT_DATE': ['дата подписания договора', report['date']],
        'REPORT_EMPLOYEE': ['ФИО сотрудника', employee['name']],
    }
    if request.method == 'POST':
        # создание нового документа
        result_params =  request.form.to_dict()
        create_report(id, result_params, goods)
        return redirect(f'/report/{id}')
    # скачивание файла, если он заполнен
    filename = f"Отчёт {report['number']} от {report['date']}.docx"
    if os.path.exists(os.path.join('reports', filename)):
        return send_from_directory(app.config['UPLOAD_FOLDER_REPORTS'], filename, as_attachment=True)
    else:
        # отрисовка формы заполнения
        flash('Договор не сформирован, заполните его')
        return render_template('report/generate_report.html', 
                               contract=contract,
                               employee=employee,
                               report_params=report_params,
                               auto_params=report_params_auto)


@app.route("/delete_report/<int:report_id>", methods=('GET', "POST", "DELETE"))
def delete_report(report_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """DELETE FROM reports WHERE id = ?""", 
        (report_id,)
    )
    conn.commit()
    conn.close()
    return redirect("/reports")


#python-docx
def create_report(id, report_params, goods):
    """ Создание нового отчёта по шаблону """
    template = os.path.join('reports', 'report_template.docx')
    result = os.path.join('reports', f"Отчёт {report_params['REPORT_NUMBER']} от {report_params['REPORT_DATE']}.docx")
    template_doc = Document(template)
    
    count = 0
    for i in range(len(goods)):
        row_cells = template_doc.tables[0].add_row().cells
        row_cells[0].text = str(i+1)
        row_cells[1].text = str(goods[0])
        row_cells[2].text = str(goods[1])
        row_cells[3].text = str(goods[2])
        count += 1
    
    report_params["REPORT_COUNT"] = count
    
    for key, value in report_params.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, "{{"+f'{str(key)}'+"}}", value)
    
        
    template_doc.save(result)


###############!            ###############
###############! Сотрудники ###############
###############!            ###############

@app.route('/employees')
def employees():
    """ Страница-список - получение всех сотрудников """
    cur = get_db_connection()
    query_result = cur.execute("""
        SELECT * FROM employees"""
    ).fetchall()
    cur.close()
    return render_template('employee/employees.html', employees=query_result)


# используется в employee(id)
def get_employee(item_id) -> sqlite3.Cursor:
    """ Получение одного сотрудника из БД """
    conn = get_db_connection()
    item: sqlite3.Cursor = conn.execute(
            """SELECT * FROM employees 
            WHERE id = ?""", 
            (item_id,)
        ).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/employee/<int:employee_id>')
def employee(employee_id):
    """ Страница-карточка - 1 контракт """
    pos = get_employee(employee_id)
    chief_id = pos['chief_id']
    if chief_id != 0:
        chief = get_employee(chief_id)
        return render_template('employee/employee.html', employee=pos,
                               chief_employee=chief)
    else:
        return render_template('employee/employee.html', employee=pos,
                               chief_employee="")


@app.route('/new_employee', methods=('GET', 'POST'))
def new_employee():
    if request.method == 'POST':
        conn = get_db_connection()
        res = conn.execute("""SELECT chief_id FROM employees""").fetchall()
        conn.close()
        chiefs = []
        for r in res:
            chiefs.append(int(*r))
        try:
            name = request.form['name'] # text
            email = request.form['email'] # text
            phone_number = request.form['phone_number'] # text
            position = request.form['position'] # text
            department = request.form['department'] # text
            chief_id = int(request.form.get('chief')) # int  
            #get shit together yo
            bundle = [
                name, email, phone_number, 
                position, department, chief_id
            ]
        except ValueError:
            # name, email, phone_number, position, department, chief_id = "", "", "", "", "", -1
            flash('Некорректные значения')
        if not all((name, email, phone_number, 
                position, department)):
            flash('Не все поля заполнены')
        elif chief_id == 0 and (0 in chiefs):
            flash('В компании уже есть руководитель')
        else:    
            conn = get_db_connection()
            cur = conn.cursor()
            cur.execute(
                inserts["employees"],
                bundle)
            conn.commit()
            employee_id = cur.lastrowid
            cur.close()
            return redirect(f'/employee/{employee_id}')
    conn = get_db_connection()
    res = conn.execute("""SELECT * FROM 'employees'""").fetchall()
    conn.close()
    return render_template('employee/new_employee.html',
                           employees=res)


@app.route('/edit_employee', methods=('GET', 'POST'))
def edit_employee():
    if request.method == 'POST':
        conn = get_db_connection()
        res = conn.execute("""SELECT chief_id FROM employees""").fetchall()
        conn.close()
        chiefs = []
        for r in res:
            chiefs.append(int(*r))
        try:
            name = request.form['name'] # text
            email = request.form['email'] # text
            phone_number = request.form['phone_number'] # text
            position = request.form['position'] # text
            department = request.form['department'] # text
            chief_id = int(request.form.get('chief')) # int  
            #get shit together yo
            bundle = [
                name, email, phone_number, 
                position, department, chief_id
            ]
        except ValueError:
            # name, email, phone_number, position, department, chief_id = "", "", "", "", "", -1
            flash('Некорректные значения')
        if not all((name, email, phone_number, 
                position, department)):
            flash('Не все поля заполнены')
        elif chief_id == 0 and (0 in chiefs):
            flash('В компании уже есть руководитель')
        else:    
            employee_id = request.args['employee_id']
            conn = get_db_connection()
            cur = conn.cursor()
            buffer = bundle[::]
            buffer.append(employee_id)
            cur.execute(
                updates["employees"],
                buffer)
            conn.commit()
            cur.close()
            return redirect(f'/employee/{employee_id}')
    conn = get_db_connection()
    employee = get_employee(request.args['employee_id'])
    res = conn.execute("""SELECT * FROM 'employees'""").fetchall()
    conn.close()
    return render_template('employee/edit_employee.html',
                           employees=res, employee=employee)


@app.route("/delete_employee/<int:employee_id>", methods=('GET', "POST", "DELETE"))
def delete_employee(employee_id: int):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        """DELETE FROM employees WHERE id = ?""", 
        (employee_id,)
    )
    conn.commit()
    conn.close()
    return redirect("/employees")


###############!     ###############
###############! 404 ###############
###############!     ###############


@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


if __name__ == '__main__':
    app.run(debug=True, port=8888)
